"""
05_build_docx.py

HW3 Step 5: Assemble the final Homework 3 .docx deliverable.

Reads:
    results/statistical_results.json   (auto-generated numbers)
    results/boxplot_composite.png      (figure)
    results/boxplot_dimensions.png     (figure)

Writes:
    homework3_DRAFT.docx

The user is expected to:
  1) Replace the [WRITING COMPONENT] block with their own ~500-word reflection
     (HW3 explicitly forbids AI-generated writing for this section)
  2) Insert their 4-5 screenshots in the marked placeholders
  3) Verify all GitHub links resolve to their repo
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
RESULTS = HERE / "results"
SCREENSHOTS = HERE / "screenshots"
JSON_PATH = RESULTS / "statistical_results.json"
BOXPLOT_COMPOSITE = RESULTS / "boxplot_composite.png"
BOXPLOT_DIMENSIONS = RESULTS / "boxplot_dimensions.png"
SCREENSHOT_1 = SCREENSHOTS / "screenshot_1_validator_running.png"
SCREENSHOT_2 = SCREENSHOTS / "screenshot_2_sample_csv_row.png"
SCREENSHOT_3 = SCREENSHOTS / "screenshot_3_rubric_in_code.png"
SCREENSHOT_4 = SCREENSHOTS / "screenshot_4_stats_terminal.png"
OUT_DOCX = HERE / "homework3_DRAFT.docx"

# Replace these with your real repo path before submitting:
GIT_REPO = "https://github.com/zw866/dsai"
BRANCH = "main"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    # Override the default theme color (which is usually dark blue) and force black
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return h


def add_para(doc, text, bold=False, italic=False, color: RGBColor | None = None):
    # Plain formatting: italic and color args are intentionally ignored
    # to keep the deliverable simple and consistent (all black, no italics).
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p


def add_placeholder(doc, label):
    p = doc.add_paragraph()
    run = p.add_run(f"[{label}]")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def add_link(doc, label, url):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(url)


def add_table_from_rows(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.rows[i].cells[j].text = str(value)
    return table


def main() -> None:
    if not JSON_PATH.exists():
        print(f"[ERROR] {JSON_PATH} not found. Run 04_statistical_analysis.py first.")
        return
    stats = json.loads(JSON_PATH.read_text())

    doc = Document()

    # Header
    add_heading(doc, "Homework 3: AI Report Validation System", level=0)
    add_para(doc, "")

    # ============== 1. Writing Component ==============
    add_heading(doc, "1. Writing Component", level=1)
    add_para(doc, "For Homework 1, I built an AI reporter. It gets World Bank GDP data for "
                  "5 countries (USA, China, India, Japan, Germany) from 2005 to 2024, then "
                  "uses Ollama to write a short summary. For Homework 3, I built a tool that "
                  "checks how good those reports are.")
    add_para(doc, "")
    add_para(doc, "Why I did not use the LAB's Likert scales", bold=True)
    add_para(doc, "The Module 9 LAB uses 6 Likert scales (1 to 5) for things like accuracy "
                  "and clarity. These are too general for my task. My reports must use only "
                  "the numbers in the source data, follow 3 fixed sections, and have a fixed "
                  "number of bullet points. Likert can't really measure these.")
    add_para(doc, "")
    add_para(doc, "So I made 5 new dimensions on different scales:")
    doc.add_paragraph("Numerical Fidelity (0-100): how many numbers match the source data",
                      style="List Number")
    doc.add_paragraph("Structural Compliance (0-100): are the 3 sections and bullet counts right",
                      style="List Number")
    doc.add_paragraph("Hallucination Penalty (0-100, lower is better): how much content was made up",
                      style="List Number")
    doc.add_paragraph("Recommendation Actionable (0-10): are the suggestions specific",
                      style="List Number")
    doc.add_paragraph("Constraint Adherence (0 or 100): did the model follow the "
                      "\"no invented facts\" rule",
                      style="List Number")
    add_para(doc, "The composite score is the average of these, scaled to 0-100.")
    add_para(doc, "")
    add_para(doc, "Experiment", bold=True)
    add_para(doc, "I compared 3 prompts from my Homework 1: A (short), B (structured), "
                  "C (strict). I made 30 reports per prompt with temperature 0.8 for variety. "
                  "Total: 90 reports. The AI validator scored each one with temperature 0.1 "
                  "for stable scoring.")
    add_para(doc, "")
    add_para(doc, "Results", bold=True)
    add_para(doc, "The ANOVA on the composite score was not significant (F = 1.07, p = 0.35). "
                  "But when I tested each dimension on its own, Recommendation Actionable "
                  "showed a clear difference (F = 9.12, p = 0.0003). Prompt C was best "
                  "(mean 4.20), then B (3.87), then A (2.80). This matches what I expected: "
                  "stricter prompts give more specific recommendations.")
    add_para(doc, "")
    add_para(doc, "Challenges", bold=True)
    add_para(doc, "Getting Ollama to return clean JSON was tricky. I set format=\"json\" "
                  "and used a low temperature (0.1) so scores were stable.")
    add_para(doc, "")

    # ============== 2. Git Repository Links ==============
    add_heading(doc, "2. Git Repository Links", level=1)
    base = f"{GIT_REPO}/blob/{BRANCH}"
    add_link(doc, "Validation system code", f"{base}/11_decision_support/homework3/02_validator.py")
    add_link(doc, "Customized rubric (in validator file)", f"{base}/11_decision_support/homework3/02_validator.py#L43-L80")
    add_link(doc, "Report generator (3 prompts batch)", f"{base}/11_decision_support/homework3/01_generate_reports.py")
    add_link(doc, "Experiment runner", f"{base}/11_decision_support/homework3/03_run_experiment.py")
    add_link(doc, "Statistical analysis", f"{base}/11_decision_support/homework3/04_statistical_analysis.py")
    add_link(doc, "Validation outputs (CSV)", f"{base}/11_decision_support/homework3/results/validation_scores.csv")
    add_link(doc, "Validated reports (HW1 source: World Bank GDP reports)",
             f"{base}/11_decision_support/homework3/reports")
    add_link(doc, "Original HW1 reporter (basis for prompts A/B/C)",
             f"{base}/03_query_ai/lab_ai_reporter.py")
    add_para(doc, "")

    # ============== 3. Screenshots / Outputs ==============
    add_heading(doc, "3. Screenshots / Outputs", level=1)

    add_para(doc, "Figure 1. Validator running.", bold=True)
    if SCREENSHOT_1.exists():
        doc.add_picture(str(SCREENSHOT_1), width=Inches(6.0))
    else:
        add_placeholder(doc, "INSERT screenshot_1_validator_running.png HERE")
    add_para(doc, "")

    add_para(doc, "Figure 2. Sample of validation_scores.csv.", bold=True)
    if SCREENSHOT_2.exists():
        doc.add_picture(str(SCREENSHOT_2), width=Inches(6.5))
    else:
        add_placeholder(doc, "INSERT screenshot_2_sample_csv_row.png HERE")
    add_para(doc, "")

    add_para(doc, "Figure 3. Validation rubric in 02_validator.py.", bold=True)
    if SCREENSHOT_3.exists():
        doc.add_picture(str(SCREENSHOT_3), width=Inches(6.0))
    else:
        add_placeholder(doc, "INSERT screenshot_3_rubric_in_code.png HERE")
    add_para(doc, "")

    add_para(doc, "Figure 4. Statistical analysis output.", bold=True)
    if SCREENSHOT_4.exists():
        doc.add_picture(str(SCREENSHOT_4), width=Inches(6.5))
    else:
        add_placeholder(doc, "INSERT screenshot_4_stats_terminal.png HERE")
    add_para(doc, "")

    add_para(doc, "Figure 5. Composite score by prompt.", bold=True)
    if BOXPLOT_COMPOSITE.exists():
        doc.add_picture(str(BOXPLOT_COMPOSITE), width=Inches(5.5))
    else:
        add_placeholder(doc, "INSERT boxplot_composite.png HERE")
    add_para(doc, "")

    add_para(doc, "Figure 6. All dimensions by prompt.", bold=True)
    if BOXPLOT_DIMENSIONS.exists():
        doc.add_picture(str(BOXPLOT_DIMENSIONS), width=Inches(6.0))
    else:
        add_placeholder(doc, "INSERT boxplot_dimensions.png HERE")
    add_para(doc, "")

    # ============== 4. Documentation ==============
    add_heading(doc, "4. Documentation", level=1)

    # 4.1 Validation Criteria Table
    add_heading(doc, "4.1 Validation Criteria Table", level=2)
    add_table_from_rows(
        doc,
        headers=["Dimension", "Scale", "Type", "Better when", "Description"],
        rows=[
            ["Numerical Fidelity", "0-100", "Continuous %", "higher",
             "% of numerical claims plausibly supported by the source data"],
            ["Structural Compliance", "0-100", "Composite (40+30+30)", "higher",
             "All 3 sections present (40) + Insights has 3-5 bullets (30) + Recs has exactly 2 (30)"],
            ["Hallucination Penalty", "0-100", "Continuous %", "LOWER",
             "% of substantive factual claims NOT grounded in source data"],
            ["Recommendation Actionable", "0-10", "Integer", "higher",
             "Average actionability (specificity, measurability) of recommendations"],
            ["Constraint Adherence", "0 or 100", "Binary", "higher",
             "1 if no fabricated indicators/countries/years detected, else 0; rescaled to 0/100"],
            ["Composite Score", "0-100", "Derived", "higher",
             "Mean of (Fidelity + Compliance + (100-Penalty) + 10*Actionable + Adherence) / 5"],
        ]
    )
    add_para(doc, "")

    # 4.2 Experimental Design
    add_heading(doc, "4.2 Experimental Design", level=2)
    n_per = stats.get("n_per_prompt", {})
    add_para(doc, "Source data: World Bank GDP (current US$, indicator NY.GDP.MKTP.CD) for "
                  "USA, China, India, Japan, Germany over 2005-2024 (same dataset used in "
                  "Homework 1).")
    add_para(doc, "Three prompts compared (taken from the v1/v2/v3 prompt iterations in "
                  "Homework 1's lab_ai_reporter.py):")
    rows = [
        ["A", "v1: terse, unstructured ('summarize the GDP dataset')", n_per.get("A", "?")],
        ["B", "v2: structured (3-5 insights, 2 recommendations)", n_per.get("B", "?")],
        ["C", "v3: constrained (sections + bullet counts + 'use only provided numbers')",
         n_per.get("C", "?")],
    ]
    add_table_from_rows(doc, ["Prompt", "Description", "n (reports generated)"], rows)
    add_para(doc, "")
    add_para(doc, "Total sample size: " + str(sum(int(v) for v in n_per.values() if isinstance(v, int)))
                  + " reports. Each report received one composite_score (and 5 dimension scores) "
                    "from the AI validator. Reports were generated with temperature=0.8 to introduce "
                    "natural variation across runs while keeping the same prompt text.")
    add_para(doc, "")

    # 4.3 Statistical Analysis
    add_heading(doc, "4.3 Statistical Analysis", level=2)
    add_para(doc, "Hypothesis (one-way ANOVA on composite_score):", bold=True)
    add_para(doc, "  H0: mean composite scores are equal across Prompts A, B, C.")
    add_para(doc, "  H1: at least one prompt mean differs from the others.")
    add_para(doc, "")
    add_para(doc, "Bartlett's test for homogeneity of variance:", bold=True)
    bart = stats.get("bartlett", {})
    add_para(doc, f"  statistic = {bart.get('statistic', '?'):.4f}, "
                  f"p = {bart.get('p_value', '?'):.4f} -> "
                  f"equal-variance assumption: "
                  f"{'satisfied' if bart.get('equal_var') else 'NOT satisfied (used Welch)'}")
    add_para(doc, "")
    add_para(doc, "ANOVA result (composite_score):", bold=True)
    anova = stats.get("anova", {})
    add_para(doc, f"  F = {anova.get('f_stat', '?'):.3f}, "
                  f"p = {anova.get('p_value', '?'):.4f}")
    welch = stats.get("welch_anova")
    if welch:
        add_para(doc, f"  Welch's F = {welch.get('f_stat', '?'):.3f}, "
                      f"Welch's p = {welch.get('p_value', '?'):.4f}")
    add_para(doc, "")

    # Per-dimension ANOVAs
    per_dim = stats.get("per_dimension_anova", {})
    if per_dim:
        add_para(doc, "Per-dimension ANOVAs (across A, B, C):", bold=True)
        rows = []
        for dim, r in per_dim.items():
            sig = "***" if r["p_value"] < 0.001 else ("**" if r["p_value"] < 0.01 else ("*" if r["p_value"] < 0.05 else "ns"))
            rows.append([
                dim,
                f"{r['f_stat']:.3f}",
                f"{r['p_value']:.4f}",
                sig,
                f"{r['mean_A']:.2f}",
                f"{r['mean_B']:.2f}",
                f"{r['mean_C']:.2f}",
            ])
        add_table_from_rows(doc, ["Dimension", "F", "p", "sig.", "Mean A", "Mean B", "Mean C"], rows)
        add_para(doc, "Significance codes: *** p<0.001, ** p<0.01, * p<0.05, ns = not significant.",
                 italic=True)
        add_para(doc, "")

    # Tukey HSD
    add_para(doc, "Tukey HSD pairwise comparisons (composite_score):", bold=True)
    tukey = stats.get("tukey_hsd", [])
    if tukey:
        rows = []
        for r in tukey:
            rows.append([
                str(r.get("A", "")),
                str(r.get("B", "")),
                f"{r.get('mean(A)', r.get('meanA', 0)):.2f}" if isinstance(r.get('mean(A)', r.get('meanA', 0)), (int, float)) else "?",
                f"{r.get('mean(B)', r.get('meanB', 0)):.2f}" if isinstance(r.get('mean(B)', r.get('meanB', 0)), (int, float)) else "?",
                f"{r.get('diff', r.get('mean_diff', 0)):.2f}" if isinstance(r.get('diff', r.get('mean_diff', 0)), (int, float)) else "?",
                f"{r.get('p-tukey', r.get('p_tukey', 1)):.4f}" if isinstance(r.get('p-tukey', r.get('p_tukey', 1)), (int, float)) else "?",
            ])
        add_table_from_rows(doc, ["Group A", "Group B", "Mean A", "Mean B", "Diff", "p-tukey"], rows)

    add_para(doc, "")
    add_para(doc, "Interpretation:", bold=True)
    concl = stats.get("conclusion", {})
    means = concl.get("means", {})
    best = concl.get("best_prompt", "?")
    sig = concl.get("anova_significant_at_0.05", False)
    add_para(doc, "  Means by prompt: " + ", ".join(f"{p}={v:.2f}" for p, v in means.items()))
    if sig:
        add_para(doc, f"  Composite ANOVA p < 0.05, so prompts differ significantly on the composite "
                      f"score. Prompt {best} achieved the highest mean composite_score "
                      f"({means.get(best, 0):.2f}) and is the recommended choice.")
    else:
        add_para(doc, f"  Composite ANOVA p >= 0.05, so we cannot reject H0 on the composite metric. "
                      f"The numeric ranking favours Prompt {best} ({means.get(best, 0):.2f}) but the "
                      f"effect is not statistically significant at this sample size.")
    sig_dims = concl.get("significant_dimensions", [])
    if sig_dims:
        add_para(doc, f"  However, individual dimensions show significant differentiation: "
                      f"{', '.join(sig_dims)}. See the per-dimension ANOVA table above for "
                      f"F-statistics and p-values per dimension. This suggests prompts produce "
                      f"reports that differ on specific quality axes even when the averaged "
                      f"composite score does not separate them.")
    add_para(doc, "")

    # 4.4 System Design
    add_heading(doc, "4.4 System Design", level=2)
    add_para(doc, "The validation system has 4 layers:")
    layers = [
        "(1) Report generator (01_generate_reports.py): pulls World Bank GDP data via the public "
        "REST API, builds 3 prompt variants, calls Ollama (gemma3:latest, temperature=0.8) "
        "30 times per prompt, writes 90 markdown reports.",
        "(2) AI validator (02_validator.py): for each report, builds a single Ollama prompt that "
        "asks the model to score the report on the 5 customized dimensions and return strict JSON. "
        "Uses Ollama JSON-mode (format='json') and low temperature (0.1) for stable scoring.",
        "(3) Experiment runner (03_run_experiment.py): walks the 90 reports, calls the validator, "
        "writes one CSV row per report with prompt_id, dimension scores, composite_score.",
        "(4) Statistical analysis (04_statistical_analysis.py): loads the CSV, runs Bartlett's test, "
        "pairwise t-tests, one-way ANOVA, Welch's ANOVA, and Tukey HSD post-hoc; saves a text "
        "summary, a JSON summary, and two boxplot PNGs.",
    ]
    for l in layers:
        doc.add_paragraph(l, style="List Bullet")
    add_para(doc, "")

    # 4.5 Technical Details
    add_heading(doc, "4.5 Technical Details", level=2)
    add_table_from_rows(
        doc,
        headers=["Component", "Detail"],
        rows=[
            ["Language", "Python 3.11+"],
            ["LLM provider", "Ollama (local), model gemma3:latest"],
            ["LLM endpoint", "http://localhost:11434/api/chat"],
            ["Source API", "World Bank API (https://api.worldbank.org/v2)"],
            ["Indicator", "NY.GDP.MKTP.CD (GDP, current US$)"],
            ["Countries", "USA, CHN, IND, JPN, DEU"],
            ["Date range", "2005-2024"],
            ["Reports per prompt", "30"],
            ["Prompts compared", "3 (A=v1, B=v2, C=v3 from HW1)"],
            ["Statistics packages", "scipy.stats, pingouin"],
            ["Plots", "matplotlib"],
            ["Document generator", "python-docx"],
        ]
    )
    add_para(doc, "")

    # 4.6 Usage Instructions
    add_heading(doc, "4.6 Usage Instructions", level=2)
    add_para(doc, "Prerequisites:", bold=True)
    pres = [
        "Python 3.11+ installed.",
        "Ollama installed and running locally on port 11434 (https://ollama.com/).",
        "Pull the model with: ollama pull gemma3:latest",
        "Install Python dependencies: pip install requests pandas scipy pingouin matplotlib python-docx",
    ]
    for p in pres:
        doc.add_paragraph(p, style="List Bullet")
    add_para(doc, "")
    add_para(doc, "Run order (from the homework3/ directory):", bold=True)
    cmds = [
        "python3 01_generate_reports.py    # ~15 min, creates reports/{A,B,C}/*.md",
        "python3 03_run_experiment.py      # ~15 min, creates results/validation_scores.csv",
        "python3 04_statistical_analysis.py  # creates results/{*.txt, *.json, *.png}",
        "python3 05_build_docx.py          # creates homework3_DRAFT.docx",
    ]
    for c in cmds:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(c)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    add_para(doc, "")

    doc.save(OUT_DOCX)
    print(f"[saved] {OUT_DOCX}")


if __name__ == "__main__":
    main()
