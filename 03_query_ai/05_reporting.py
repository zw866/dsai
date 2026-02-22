# 05_reporting.py
# Save AI Report in Multiple Formats
# Pairs with 05_reporting.R
# Tim Fraser

# This script demonstrates how to save AI-generated reports in different formats:
# .txt, .md, .html, and .docx. Students will learn how to format and write
# LLM output to various file types for different use cases.

# 0. SETUP ###################################

## 0.1 Load Packages #################################

# If you haven't already, install required packages:
#   pip install markdown python-docx
# (Use "python-docx", not "docx" — the latter is a different, broken package.)
import markdown  # for converting markdown to HTML
from docx import Document  # for creating Word documents (from python-docx)

## 0.2 Mock LLM Output #########################

# Simulate an AI response object
# In a real script, this would come from your LLM API call
mock_llm_response = {
    "response": """# Data Analysis Report

## Summary
The dataset contains 150 records with 3 key metrics showing positive trends.

## Key Findings
- Metric A increased by 15% over the period
- Metric B remained stable at 42 units
- Metric C showed significant variation

## Recommendations
Consider further investigation into Metric C variations."""
}

# Extract the text content
report_text = mock_llm_response["response"]

# 1. SAVE AS PLAIN TEXT (.txt) ###################################

# Write directly to a text file
# Simple and universal format
with open("03_query_ai/05_reporting_report.txt", "w", encoding="utf-8") as f:
    f.write(report_text)

print("✅ Saved 03_query_ai/05_reporting_report.txt")

# 2. SAVE AS MARKDOWN (.md) ###################################

# Markdown files are great for GitHub and documentation
# The content is already in markdown format, so we just write it
with open("03_query_ai/05_reporting_report.md", "w", encoding="utf-8") as f:
    f.write(report_text)

print("✅ Saved 05_reporting_report.md")

# 3. SAVE AS HTML (.html) ###################################

# Convert markdown to HTML for web-friendly viewing
html_content = markdown.markdown(report_text)

# Wrap in basic HTML structure
html_document = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Data Analysis Report</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
        h1 {{ color: #333; }}
        h2 {{ color: #666; margin-top: 30px; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

with open("03_query_ai/05_reporting_report.html", "w", encoding="utf-8") as f:
    f.write(html_document)

print("✅ Saved 05_reporting_report.html")
    
# 4. SAVE AS WORD DOCUMENT (.docx) ###################################

# Create a Word document for professional sharing
doc = Document()

# Split content by lines and add to document
# Handle markdown headers and formatting
for line in report_text.split("\n"):
    if line.startswith("# "):
        # Main heading
        doc.add_heading(line[2:], level=1)
    elif line.startswith("## "):
        # Subheading
        doc.add_heading(line[3:], level=2)
    elif line.startswith("- "):
        # Bullet point
        doc.add_paragraph(line[2:], style="List Bullet")
    elif line.strip():
        # Regular paragraph
        doc.add_paragraph(line)

doc.save("03_query_ai/05_reporting_report.docx")

print("✅ Saved 05_reporting_report.docx")
print("\n✅ All report formats saved successfully!")
